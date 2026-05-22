"""Reduce a `.docx` file to a deterministic semantic summary.

python-docx's `.docx` output is not byte-stable across versions (zip
metadata, internal id assignment, XML attribute ordering all drift),
so we never snapshot the raw bytes. Instead, summarise the document
structure into a plain `dict` that captures what matters for the
exporter contract: heading levels and text, paragraph styles, table
shapes and cell text, embedded image counts, and the set of styles
the document references.

Two callers:

- Snapshot tests: ``summarise_path(p)`` for files on disk.
- The bytes-on-stdout test for ``export docx -o -``: ``summarise_bytes(b)``.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

import docx
from docx.document import Document as DocxDocument


def summarise_path(path: Path) -> dict[str, Any]:
    """Summarise a .docx on disk."""
    doc = docx.Document(str(path))
    return _summarise(doc)


def summarise_bytes(data: bytes) -> dict[str, Any]:
    """Summarise a .docx from in-memory bytes."""
    doc = docx.Document(io.BytesIO(data))
    return _summarise(doc)


def _summarise(doc: DocxDocument) -> dict[str, Any]:
    paragraphs: list[dict[str, Any]] = []
    headings: list[tuple[int, str]] = []
    styles_used: set[str] = set()

    for p in doc.paragraphs:
        style = p.style.name if p.style is not None else ""
        styles_used.add(style)
        text = "".join(r.text for r in p.runs).strip()
        if style.startswith("Heading "):
            try:
                level = int(style.split(" ", 1)[1])
            except ValueError:
                level = 0
            headings.append((level, text))
        paragraphs.append(
            {
                "style": style,
                "text": text,
                "runs": [
                    {
                        "text": r.text,
                        "bold": bool(r.bold),
                        "italic": bool(r.italic),
                    }
                    for r in p.runs
                    if r.text
                ],
            }
        )

    tables: list[dict[str, Any]] = []
    for t in doc.tables:
        rows = []
        for row in t.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"rows": rows, "n_rows": len(rows), "n_cols": len(rows[0]) if rows else 0})

    # Image count: walk inline shapes (the only kind python-docx adds).
    images = len(doc.inline_shapes)

    return {
        "paragraphs": paragraphs,
        "headings": headings,
        "tables": tables,
        "images": images,
        "styles_used": sorted(styles_used),
    }
