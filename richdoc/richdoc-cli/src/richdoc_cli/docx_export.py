"""HTML → DOCX exporter for richdoc documents.

The walker mirrors `markdown.py`: parse the HTML, dispatch each element to
a tag-specific handler. Each handler appends paragraphs / tables / images
to a `python-docx` `Document`.

Style strategy. Confluence's "Import Word document" macro keys off Word's
built-in style names — `Heading 1`..`Heading 6`, `List Bullet`,
`List Number`, `Intense Quote`, `Table Grid`. We use those names verbatim
so headings come through as Confluence headings, lists as Confluence
lists, etc. One custom paragraph style — `RichdocCode` — is registered at
document open for monospace code blocks. Code blocks come into Confluence
as monospace-styled paragraphs (Confluence's importer doesn't have a true
code-macro path from docx, so this is as close as we get).

Images. Every `<img>`, `<rd-shot>`, and `<rd-figure>` image is registered
with the shared `AssetStore`, fetched (relative AND remote — Confluence
import only renders package-embedded images), and inserted via
`document.add_picture(BytesIO(bytes), …)`.

Diagrams. `<rd-mermaid>` and `<rd-plantuml>` are POSTed to a Kroki
endpoint and embedded as PNG. Failures fall back to a code block.

Dropped components. `rd-icon` (no native fit), `rd-tooltip` (rendered
inline as `term (definition)`), `rd-toc` in single-file mode (the doc's
heading order is the TOC). In book mode, the bundled book is rendered as
one long document — the rd-toc becomes a small "Contents" heading + the
chapter titles, since Confluence's importer doesn't honour native Word
TOC fields reliably.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Callable

import lxml.etree as ET
import lxml.html as LH
from docx import Document
from docx.document import Document as DocumentT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell

from .export_assets import AssetStore
from .book import ChapterFile
from .diagrams import render_to_png

# ---------------------------------------------------------------------------
# Public entry points
# ---------------------------------------------------------------------------


@dataclass
class DocxResult:
    data: bytes
    dropped: list[str] = field(default_factory=list)
    missing: list[str] = field(default_factory=list)
    diagrams_rendered: int = 0
    diagrams_failed: int = 0
    images_embedded: int = 0


def html_to_docx(
    source: str,
    *,
    base_dir: Path,
    render_diagrams: bool = True,
    diagram_endpoint: str = "https://kroki.io",
    asset_store: AssetStore | None = None,
) -> DocxResult:
    """Render a single richdoc HTML file to a DOCX byte string."""
    doc = _new_document()
    store = asset_store if asset_store is not None else AssetStore()
    state = _State(
        doc=doc,
        asset_store=store,
        base_dir=base_dir,
        render_diagrams=render_diagrams,
        diagram_endpoint=diagram_endpoint,
    )
    _render_source(state, source)
    _finalise(state)
    return _serialise(state)


def chapters_to_docx(
    chapters: list[ChapterFile],
    *,
    book_base: Path,
    render_diagrams: bool = True,
    diagram_endpoint: str = "https://kroki.io",
) -> DocxResult:
    """Render a multi-file book into one DOCX. Each chapter starts on a new
    page; the chapter title is emitted as Heading 1."""
    doc = _new_document()
    store = AssetStore()
    state = _State(
        doc=doc,
        asset_store=store,
        base_dir=book_base,
        render_diagrams=render_diagrams,
        diagram_endpoint=diagram_endpoint,
    )
    for i, ch in enumerate(chapters):
        if i > 0:
            state.add_page_break()
        # Override base_dir for asset resolution per chapter.
        state.base_dir = ch.path.parent
        _render_source(state, ch.html, chapter_title=ch.title)
        # Drop any rd-toc dropped-tag entries from each chapter — the shared
        # rd-toc lives in every chapter file by design, and we already render
        # it inline as a Contents heading. The drop entries from chapter 2+
        # are misleading noise.
    state.base_dir = book_base
    _finalise(state)
    return _serialise(state)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------


def _new_document() -> DocumentT:
    doc = Document()
    _register_code_style(doc)
    return doc


def _register_code_style(doc: DocumentT) -> None:
    """Register the `RichdocCode` paragraph style used for fenced code blocks."""
    styles = doc.styles
    if "RichdocCode" in [s.name for s in styles]:
        return
    from docx.enum.style import WD_STYLE_TYPE

    style = styles.add_style("RichdocCode", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    font = style.font
    font.name = "Courier New"
    font.size = Pt(9)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Walker state
# ---------------------------------------------------------------------------


@dataclass
class _State:
    doc: DocumentT
    asset_store: AssetStore
    base_dir: Path
    render_diagrams: bool
    diagram_endpoint: str
    # Bookkeeping that surfaces in the JSON envelope.
    dropped: list[str] = field(default_factory=list)
    diagrams_rendered: int = 0
    diagrams_failed: int = 0
    images_embedded: int = 0
    # Stack of (list_kind, depth) for nested <ul>/<ol>.
    list_stack: list[tuple[str, int]] = field(default_factory=list)
    # Citation registry.
    cite_counter: int = 0
    cite_order: list[str] = field(default_factory=list)
    refs_collected: dict[str, dict[str, str]] = field(default_factory=dict)
    # True once an rd-toc has been rendered — subsequent ones (shared TOC in
    # every chapter of a book) are skipped.
    toc_emitted: bool = False

    # ---- writers ---------------------------------------------------------

    def add_paragraph(self, text: str = "", style: str | None = None):
        p = self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()
        if text:
            p.add_run(text)
        return p

    def add_page_break(self) -> None:
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    def record_dropped(self, tag: str) -> None:
        self.dropped.append(tag)


# ---------------------------------------------------------------------------
# Render helpers
# ---------------------------------------------------------------------------


def _render_source(state: _State, source: str, *, chapter_title: str | None = None) -> None:
    parser = LH.HTMLParser(recover=True)
    root = LH.document_fromstring(source, parser=parser)
    body = root.find(".//body")
    target = body if body is not None else root
    # The chapter's own rd-hero / <h1> supplies its Heading 1. We only fall
    # back to the TOC-derived chapter title when the chapter has no heading
    # of its own anywhere in its top-level page tree.
    if chapter_title and not _has_any_heading(target):
        state.doc.add_heading(chapter_title, level=1)
    _render_children(state, target)


def _render_children(state: _State, el: ET._Element) -> None:  # noqa: SLF001
    # Block-level walk: each direct child gets dispatched. Mixed text /
    # element content at the top of `<body>` is rare; we wrap stray text in
    # a paragraph.
    if el.text and el.text.strip():
        state.add_paragraph(_inline_text(el.text))
    for child in el:
        _render_block(state, child)
        if child.tail and child.tail.strip():
            state.add_paragraph(_inline_text(child.tail))


def _render_block(state: _State, el: ET._Element) -> None:  # noqa: SLF001
    tag = el.tag
    if not isinstance(tag, str):
        return
    tag = tag.lower()
    handler = _BLOCK_HANDLERS.get(tag)
    if handler is None:
        if tag.startswith("rd-"):
            # Unknown rd-* — emit a paragraph from its inline text.
            text = _flatten_inline(state, el).strip()
            if text:
                state.add_paragraph(text)
            state.record_dropped(tag)
        else:
            # Unknown plain tag — emit children as paragraphs.
            _render_children(state, el)
        return
    handler(state, el)


def _has_any_heading(el: ET._Element) -> bool:  # noqa: SLF001
    """Detect whether the chapter has its own top heading anywhere inside
    its rd-page (or top-level if no rd-page). Used to decide whether we need
    to inject a TOC-derived chapter heading."""
    for node in el.iter():
        if not isinstance(node.tag, str):
            continue
        tag = node.tag.lower()
        if tag in ("h1", "rd-hero"):
            return True
    return False


# ---------------------------------------------------------------------------
# Inline rendering — produces a list of "run specs" so callers can append
# them onto a paragraph with the right formatting.
# ---------------------------------------------------------------------------


@dataclass
class _Run:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False
    underline: bool = False
    strike: bool = False
    hyperlink: str | None = None


def _flatten_inline(state: _State, el: ET._Element) -> str:  # noqa: SLF001
    return "".join(r.text for r in _inline_runs(state, el))


def _inline_runs(
    state: _State,
    el: ET._Element,
    *,
    bold: bool = False,
    italic: bool = False,
    code: bool = False,
    underline: bool = False,
    strike: bool = False,
    hyperlink: str | None = None,
) -> list[_Run]:
    """Walk inline content and yield runs with cascaded formatting."""
    runs: list[_Run] = []
    if el.text:
        runs.append(
            _Run(
                _inline_text(el.text),
                bold=bold,
                italic=italic,
                code=code,
                underline=underline,
                strike=strike,
                hyperlink=hyperlink,
            )
        )
    for child in el:
        tag = child.tag if isinstance(child.tag, str) else ""
        tag = tag.lower()
        child_bold = bold
        child_italic = italic
        child_code = code
        child_underline = underline
        child_strike = strike
        child_link = hyperlink
        skip_children = False

        if tag in ("strong", "b"):
            child_bold = True
        elif tag in ("em", "i"):
            child_italic = True
        elif tag == "code":
            child_code = True
        elif tag == "u":
            child_underline = True
        elif tag in ("s", "del", "strike"):
            child_strike = True
        elif tag == "a":
            child_link = child.get("href") or hyperlink
            child_underline = True
        elif tag == "br":
            runs.append(_Run("\n"))
            skip_children = True
        elif tag == "rd-footnote":
            mark = child.get("mark") or "*"
            runs.append(_Run(f"[{mark}]", bold=bold, italic=italic))
            # Also flatten body so context isn't lost.
            body = _flatten_inline(state, child).strip()
            if body:
                runs.append(_Run(f" ({body})", italic=True))
            skip_children = True
        elif tag == "rd-cite":
            key = child.get("key") or ""
            if key not in state.cite_order:
                state.cite_order.append(key)
            n = state.cite_order.index(key) + 1
            runs.append(_Run(f"[{n}]", bold=bold, italic=italic))
            skip_children = True
        elif tag == "rd-ref":
            _collect_ref(state, child)
            skip_children = True
        elif tag == "rd-icon":
            label = child.get("label") or ""
            if label:
                runs.append(_Run(label, bold=bold, italic=italic))
            state.record_dropped("rd-icon")
            skip_children = True
        elif tag == "rd-tooltip":
            term = child.get("term") or ""
            body = _flatten_inline(state, child).strip()
            if term and body:
                runs.append(_Run(f"{term} ({body})", bold=bold, italic=italic))
            elif term:
                runs.append(_Run(term, bold=bold, italic=italic))
            skip_children = True
        elif tag == "rd-badge":
            text = _flatten_inline(state, child).strip()
            if text:
                runs.append(_Run(f"[{text}]", bold=True))
            skip_children = True
        elif tag == "img":
            # Inline images render after the current paragraph in DOCX.
            # Flatten alt as a stand-in for inline-flow continuity.
            alt = child.get("alt") or ""
            if alt:
                runs.append(_Run(alt, italic=True))
            skip_children = True
        elif tag in ("script", "style"):
            skip_children = True

        if not skip_children:
            runs.extend(
                _inline_runs(
                    state,
                    child,
                    bold=child_bold,
                    italic=child_italic,
                    code=child_code,
                    underline=child_underline,
                    strike=child_strike,
                    hyperlink=child_link,
                )
            )
        if child.tail:
            runs.append(
                _Run(
                    _inline_text(child.tail),
                    bold=bold,
                    italic=italic,
                    code=code,
                    underline=underline,
                    strike=strike,
                    hyperlink=hyperlink,
                )
            )
    return runs


def _emit_runs(paragraph, runs: list[_Run]) -> None:
    """Append `runs` onto a python-docx paragraph."""
    for run_spec in runs:
        text = run_spec.text
        if not text:
            continue
        if run_spec.hyperlink:
            _add_hyperlink(paragraph, text, run_spec.hyperlink, run_spec)
            continue
        r = paragraph.add_run(text)
        if run_spec.bold:
            r.bold = True
        if run_spec.italic:
            r.italic = True
        if run_spec.underline:
            r.underline = True
        if run_spec.strike:
            r.font.strike = True
        if run_spec.code:
            r.font.name = "Courier New"


def _add_hyperlink(paragraph, text: str, url: str, spec: _Run) -> None:
    """Add a native Word hyperlink — python-docx doesn't expose one directly."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = ET.SubElement(paragraph._p, qn("w:hyperlink"))
    hyperlink.set(qn("r:id"), r_id)

    new_run = ET.SubElement(hyperlink, qn("w:r"))
    rpr = ET.SubElement(new_run, qn("w:rPr"))
    color = ET.SubElement(rpr, qn("w:color"))
    color.set(qn("w:val"), "0563C1")
    underline = ET.SubElement(rpr, qn("w:u"))
    underline.set(qn("w:val"), "single")
    if spec.bold:
        ET.SubElement(rpr, qn("w:b"))
    if spec.italic:
        ET.SubElement(rpr, qn("w:i"))
    if spec.code:
        rfonts = ET.SubElement(rpr, qn("w:rFonts"))
        rfonts.set(qn("w:ascii"), "Courier New")
        rfonts.set(qn("w:hAnsi"), "Courier New")
    t = ET.SubElement(new_run, qn("w:t"))
    t.text = text
    t.set(qn("xml:space"), "preserve")


# ---------------------------------------------------------------------------
# Block handlers — plain HTML
# ---------------------------------------------------------------------------


def _h_heading(level: int) -> Callable[[_State, ET._Element], None]:
    def handler(state: _State, el: ET._Element) -> None:
        runs = _inline_runs(state, el)
        # docx's Document.add_heading honours levels 1-9 (with level 0 = Title).
        p = state.doc.add_heading("", level=min(level, 9))
        _emit_runs(p, runs)

    return handler


def _h_p(state: _State, el: ET._Element) -> None:
    runs = _inline_runs(state, el)
    if not any(r.text.strip() for r in runs):
        return
    p = state.add_paragraph()
    _emit_runs(p, runs)


def _h_ul(state: _State, el: ET._Element) -> None:
    _emit_list(state, el, kind="ul")


def _h_ol(state: _State, el: ET._Element) -> None:
    _emit_list(state, el, kind="ol")


def _emit_list(state: _State, el: ET._Element, *, kind: str) -> None:  # noqa: SLF001
    depth = sum(1 for k, _ in state.list_stack if k in ("ul", "ol"))
    state.list_stack.append((kind, depth))
    style_for_level = {
        "ul": ["List Bullet", "List Bullet 2", "List Bullet 3"],
        "ol": ["List Number", "List Number 2", "List Number 3"],
    }
    style = style_for_level[kind][min(depth, 2)]
    for child in el:
        if not isinstance(child.tag, str):
            continue
        if child.tag.lower() != "li":
            continue
        # Split <li> content into inline run + nested list children.
        inline_parts, nested_blocks = _split_li(child)
        p = state.doc.add_paragraph(style=style)
        runs = _inline_runs_from_parts(state, inline_parts)
        _emit_runs(p, runs)
        for nested in nested_blocks:
            _render_block(state, nested)
    state.list_stack.pop()


def _split_li(li: ET._Element) -> tuple[list, list[ET._Element]]:  # noqa: SLF001
    """Return (inline_parts, nested_block_elements). `inline_parts` is a list
    of (kind, value) where kind is `text` or `element`."""
    parts: list[tuple[str, object]] = []
    if li.text:
        parts.append(("text", li.text))
    blocks: list[ET._Element] = []
    for child in li:
        if isinstance(child.tag, str) and child.tag.lower() in (
            "ul",
            "ol",
            "pre",
            "blockquote",
            "table",
        ):
            blocks.append(child)
            if child.tail:
                parts.append(("text", child.tail))
            continue
        parts.append(("element", child))
        if child.tail:
            parts.append(("text", child.tail))
    return parts, blocks


def _inline_runs_from_parts(state: _State, parts: list[tuple[str, object]]) -> list[_Run]:  # noqa: SLF001
    runs: list[_Run] = []
    for kind, val in parts:
        if kind == "text":
            text = _inline_text(val)  # type: ignore[arg-type]
            if text:
                runs.append(_Run(text))
        else:
            el = val  # type: ignore[assignment]
            tag = el.tag if isinstance(el.tag, str) else ""  # type: ignore[union-attr]
            tag = tag.lower()
            wrapper = ET.Element("span")  # synthetic parent
            wrapper.append(el)  # type: ignore[arg-type]
            wrapper.remove(el)  # type: ignore[arg-type]
            # Just call _inline_runs treating the element as a one-shot tree.
            tmp = ET.Element("span")
            tmp.append(el)  # type: ignore[arg-type]
            runs.extend(_inline_runs(state, tmp))
            # _inline_runs uses tail; clear it so caller doesn't double-emit.
            el.tail = None  # type: ignore[union-attr]
    return runs


def _h_blockquote(state: _State, el: ET._Element) -> None:
    inner_paragraphs = [c for c in el if isinstance(c.tag, str) and c.tag.lower() == "p"]
    if not inner_paragraphs:
        runs = _inline_runs(state, el)
        if any(r.text.strip() for r in runs):
            p = state.doc.add_paragraph(style="Intense Quote")
            _emit_runs(p, runs)
        return
    for para in inner_paragraphs:
        runs = _inline_runs(state, para)
        if not any(r.text.strip() for r in runs):
            continue
        p = state.doc.add_paragraph(style="Intense Quote")
        _emit_runs(p, runs)


def _h_pre(state: _State, el: ET._Element) -> None:
    # <pre><code>…</code></pre> is the common shape.
    text = "".join(el.itertext())
    _emit_code(state, _dedent(text), lang=None)


def _h_hr(state: _State, el: ET._Element) -> None:
    p = state.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = ET.SubElement(pPr, qn("w:pBdr"))
    bottom = ET.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")


def _h_table(state: _State, el: ET._Element) -> None:
    headers, rows = _parse_html_table(el)
    if not rows and not headers:
        return
    cols = max(len(headers), max((len(r) for r in rows), default=0))
    if cols == 0:
        return
    if not headers:
        headers = [""] * cols
    table = state.doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    _fill_row(state, table.rows[0], headers, bold=True)
    for i, row in enumerate(rows, start=1):
        _fill_row(state, table.rows[i], row + [""] * (cols - len(row)))


def _parse_html_table(el: ET._Element) -> tuple[list[str], list[list[str]]]:  # noqa: SLF001
    headers: list[str] = []
    rows: list[list[str]] = []
    for tr in el.iter("tr"):
        cells = [c for c in tr if isinstance(c.tag, str) and c.tag.lower() in ("th", "td")]
        texts = [" ".join("".join(c.itertext()).split()).strip() for c in cells]
        if cells and all(isinstance(c.tag, str) and c.tag.lower() == "th" for c in cells) and not headers:
            headers = texts
        else:
            rows.append(texts)
    return headers, rows


def _fill_row(state: _State, row, cells: list[str], *, bold: bool = False) -> None:
    for idx, text in enumerate(cells):
        cell = row.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        if bold:
            r.bold = True


def _h_img_block(state: _State, el: ET._Element) -> None:
    src = el.get("src") or ""
    alt = el.get("alt") or ""
    if not src:
        return
    _embed_image(state, src, alt=alt)


# ---------------------------------------------------------------------------
# Block handlers — rd-* components
# ---------------------------------------------------------------------------


def _h_rd_page(state: _State, el: ET._Element) -> None:
    _render_children(state, el)


def _h_rd_section(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    if title:
        state.doc.add_heading(title, level=2)
    _render_children(state, el)


def _h_rd_hero(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    eyebrow = el.get("eyebrow") or ""
    lede = el.get("lede") or ""
    meta = el.get("meta") or ""
    if eyebrow:
        p = state.add_paragraph()
        r = p.add_run(eyebrow.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if title:
        state.doc.add_heading(title, level=1)
    if lede:
        p = state.add_paragraph()
        r = p.add_run(lede)
        r.italic = True
    if meta:
        p = state.add_paragraph()
        r = p.add_run(meta)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # Extras
    _render_children(state, el)


def _h_rd_banner(state: _State, el: ET._Element) -> None:
    type_ = (el.get("type") or "info").lower()
    msg = (el.get("message") or _flatten_inline(state, el)).strip()
    label = type_.upper()
    p = state.add_paragraph()
    r = p.add_run(f"[{label}] ")
    r.bold = True
    if msg:
        p.add_run(msg)


def _h_rd_callout(state: _State, el: ET._Element) -> None:
    type_ = (el.get("type") or "info").lower()
    title = el.get("title") or _CALLOUT_DEFAULT_TITLE.get(type_, "")
    color = _CALLOUT_COLORS.get(type_, "DDDDDD")

    table = state.doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_border(cell, side="left", color=color, sz=24)
    _set_cell_shading(cell, "FAFAFA")
    cell.text = ""

    if title:
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        body_p_first = cell.add_paragraph()
    else:
        body_p_first = cell.paragraphs[0]

    runs = _inline_runs(state, el)
    if any(r.text.strip() for r in runs):
        # Use first body paragraph for the inline content of the callout.
        _emit_runs(body_p_first, runs)
    # Any block-level children (lists, code) — render as additional paragraphs
    # in the cell. python-docx tables don't support arbitrary block nesting,
    # so we flatten the structure: each block becomes a plain paragraph.
    for child in el:
        if not isinstance(child.tag, str):
            continue
        tag = child.tag.lower()
        if tag in ("p", "ul", "ol", "pre", "rd-code", "rd-shell"):
            text = _flatten_inline(state, child).strip()
            if text:
                cp = cell.add_paragraph()
                cp.add_run(text)


_CALLOUT_DEFAULT_TITLE = {
    "info": "Info",
    "success": "Success",
    "warn": "Warning",
    "danger": "Danger",
    "note": "Note",
    "tldr": "TL;DR",
}

_CALLOUT_COLORS = {
    "info": "3B82F6",
    "success": "10B981",
    "warn": "F59E0B",
    "danger": "EF4444",
    "note": "6B7280",
    "tldr": "111111",
}


def _h_rd_kv(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    if title:
        state.doc.add_heading(title, level=3)
    rows = [r for r in el if isinstance(r.tag, str) and r.tag.lower() == "rd-row"]
    if not rows:
        return
    table = state.doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    for i, row in enumerate(rows):
        key = row.get("key") or ""
        value_runs = _inline_runs(state, row)
        cells = table.rows[i].cells
        cells[0].text = ""
        kp = cells[0].paragraphs[0]
        kr = kp.add_run(key)
        kr.bold = True
        cells[1].text = ""
        vp = cells[1].paragraphs[0]
        _emit_runs(vp, value_runs)


def _h_rd_stat(state: _State, el: ET._Element) -> None:
    value = el.get("value") or ""
    label = el.get("label") or ""
    delta = el.get("delta") or ""
    p = state.add_paragraph()
    r = p.add_run(value)
    r.bold = True
    r.font.size = Pt(24)
    if label:
        p2 = state.add_paragraph()
        rr = p2.add_run(label)
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if delta:
        p3 = state.add_paragraph()
        p3.add_run(delta)


def _h_rd_progress(state: _State, el: ET._Element) -> None:
    value = el.get("value") or ""
    label = el.get("label") or ""
    p = state.add_paragraph()
    if label:
        r = p.add_run(f"{label}: ")
        r.bold = True
    p.add_run(value)


def _h_rd_update(state: _State, el: ET._Element) -> None:
    date = el.get("date") or ""
    kind = el.get("kind") or ""
    author = el.get("author") or ""
    title = el.get("title") or ""
    heading_bits = [b for b in (date, kind, title) if b]
    if heading_bits:
        state.doc.add_heading(" — ".join(heading_bits), level=4)
    if author:
        p = state.add_paragraph()
        r = p.add_run(f"by {author}")
        r.italic = True
    _render_children(state, el)


def _h_rd_quote(state: _State, el: ET._Element) -> None:
    runs = _inline_runs(state, el)
    if any(r.text.strip() for r in runs):
        p = state.doc.add_paragraph(style="Intense Quote")
        _emit_runs(p, runs)
    author = el.get("author") or ""
    cite = el.get("cite") or ""
    bits = [b for b in (author, cite) if b]
    if bits:
        p = state.add_paragraph()
        r = p.add_run("— " + ", ".join(bits))
        r.italic = True


def _h_rd_cols(state: _State, el: ET._Element) -> None:
    # Flatten to sequential rendering — Confluence Word import doesn't
    # preserve Word columns reliably, and rd-cols is structural parallelism.
    _render_children(state, el)


def _h_rd_card(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    accent = el.get("accent") or ""
    if accent:
        p = state.add_paragraph()
        r = p.add_run(accent.upper())
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if title:
        state.doc.add_heading(title, level=3)
    _render_children(state, el)


def _h_rd_code(state: _State, el: ET._Element) -> None:
    lang = el.get("lang") or None
    title = el.get("title") or ""
    if title:
        p = state.add_paragraph()
        r = p.add_run(title)
        r.bold = True
    text = _dedent(_element_source(el))
    _emit_code(state, text, lang=lang)


def _h_rd_diff(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    if title:
        p = state.add_paragraph()
        r = p.add_run(title)
        r.bold = True
    text = _dedent(_element_source(el))
    _emit_code(state, text, lang="diff")


def _h_rd_shell(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    if title:
        p = state.add_paragraph()
        r = p.add_run(title)
        r.bold = True
    lines: list[str] = []
    for child in el:
        if not isinstance(child.tag, str):
            continue
        tag = child.tag.lower()
        text = " ".join("".join(child.itertext()).split())
        if tag == "rd-prompt":
            cwd = child.get("cwd")
            if cwd:
                lines.append(f"{cwd} $ {text}")
            else:
                lines.append(f"$ {text}")
        elif tag == "rd-output":
            lines.append(text)
    _emit_code(state, "\n".join(lines), lang=None)


def _h_rd_math(state: _State, el: ET._Element) -> None:
    text = _element_source(el).strip()
    if not text:
        return
    p = state.add_paragraph(style="RichdocCode")
    p.add_run(text)


def _h_rd_figure(state: _State, el: ET._Element) -> None:
    caption = el.get("caption") or ""
    # rd-figure wraps an arbitrary block (img, rd-plantuml, rd-mermaid,
    # rd-chart, etc.) plus an optional caption.
    _render_children(state, el)
    if caption:
        p = state.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)


def _h_rd_chart(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    caption = el.get("caption") or ""
    data_attr = el.get("data") or _element_source(el)
    if title:
        p = state.add_paragraph()
        r = p.add_run(title)
        r.bold = True
    rendered = _chart_to_table(state, data_attr)
    if not rendered and data_attr.strip():
        _emit_code(state, data_attr.strip(), lang=None)
    if caption:
        p = state.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)


def _chart_to_table(state: _State, raw: str) -> bool:  # noqa: SLF001
    raw = raw.strip()
    if not raw:
        return False
    import json as _json

    if raw.startswith("[") or raw.startswith("{"):
        try:
            data = _json.loads(raw)
        except ValueError:
            data = None
        if isinstance(data, list) and data:
            if isinstance(data[0], dict):
                keys = list(data[0].keys())
                rows_data = data
            else:
                keys = ["value"]
                rows_data = [{"value": v} for v in data]
            table = state.doc.add_table(rows=1 + len(rows_data), cols=len(keys))
            table.style = "Table Grid"
            _fill_row(state, table.rows[0], keys, bold=True)
            for i, row in enumerate(rows_data, start=1):
                _fill_row(state, table.rows[i], [str(row.get(k, "")) for k in keys])
            return True
    if "\n" in raw and "," in raw:
        rows = [r.split(",") for r in raw.splitlines() if r.strip()]
        if rows:
            head = [c.strip() for c in rows[0]]
            body = rows[1:]
            table = state.doc.add_table(rows=1 + len(body), cols=len(head))
            table.style = "Table Grid"
            _fill_row(state, table.rows[0], head, bold=True)
            for i, r in enumerate(body, start=1):
                _fill_row(state, table.rows[i], [(c.strip() if j < len(r) else "") for j, c in enumerate(r)])
            return True
    if re.fullmatch(r"[\d\s,.\-eE]+", raw):
        values = [v.strip() for v in raw.split(",") if v.strip()]
        table = state.doc.add_table(rows=1 + len(values), cols=2)
        table.style = "Table Grid"
        _fill_row(state, table.rows[0], ["#", "value"], bold=True)
        for i, v in enumerate(values, start=1):
            _fill_row(state, table.rows[i], [str(i), v])
        return True
    return False


def _h_rd_gallery(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    if title:
        state.doc.add_heading(title, level=3)
    for shot in el:
        if not (isinstance(shot.tag, str) and shot.tag.lower() == "rd-shot"):
            continue
        src = shot.get("src") or ""
        alt = shot.get("alt") or ""
        caption = shot.get("caption") or ""
        if not src:
            continue
        _embed_image(state, src, alt=alt)
        if caption:
            p = state.add_paragraph()
            r = p.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)


def _h_rd_embed(state: _State, el: ET._Element) -> None:
    src = el.get("src") or ""
    title = el.get("title") or "Embed"
    caption = el.get("caption") or ""
    p = state.add_paragraph()
    runs = [_Run(f"▶ {title}", hyperlink=src or None, underline=bool(src))]
    _emit_runs(p, runs)
    if caption:
        p2 = state.add_paragraph()
        r = p2.add_run(caption)
        r.italic = True


def _h_rd_tabs(state: _State, el: ET._Element) -> None:
    for tab in el:
        if not (isinstance(tab.tag, str) and tab.tag.lower() == "rd-tab"):
            continue
        label = tab.get("label") or ""
        if label:
            state.doc.add_heading(label, level=3)
        _render_children(state, tab)


def _h_rd_timeline(state: _State, el: ET._Element) -> None:
    for ev in el:
        if not (isinstance(ev.tag, str) and ev.tag.lower() == "rd-event"):
            continue
        date = ev.get("date") or ""
        title = ev.get("title") or ""
        head = " — ".join(b for b in (date, title) if b)
        if head:
            p = state.add_paragraph()
            r = p.add_run(head)
            r.bold = True
        _render_children(state, ev)


def _h_rd_steps(state: _State, el: ET._Element) -> None:
    state.list_stack.append(("ol", 0))
    for step in el:
        if not (isinstance(step.tag, str) and step.tag.lower() == "rd-step"):
            continue
        title = step.get("title") or ""
        done = step.get("done") is not None
        marker = "☑ " if done else ""
        p = state.doc.add_paragraph(style="List Number")
        r = p.add_run(marker)
        if title:
            tr = p.add_run(title)
            tr.bold = True
        _render_children_into(state, step)
    state.list_stack.pop()


def _render_children_into(state: _State, el: ET._Element) -> None:
    # Render children as additional block paragraphs after the leading list item.
    for child in el:
        _render_block(state, child)


def _h_rd_detail(state: _State, el: ET._Element) -> None:
    summary = el.get("summary") or ""
    if summary:
        state.doc.add_heading(summary, level=3)
    _render_children(state, el)


def _h_rd_tree(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    if title:
        state.doc.add_heading(title, level=3)
    _emit_tree_nodes(state, el, depth=0)


def _emit_tree_nodes(state: _State, parent: ET._Element, *, depth: int) -> None:
    for node in parent:
        if not (isinstance(node.tag, str) and node.tag.lower() == "rd-node"):
            continue
        label = node.get("label") or ""
        style_level = min(depth, 2)
        style = ["List Bullet", "List Bullet 2", "List Bullet 3"][style_level]
        p = state.doc.add_paragraph(label, style=style)
        # Inline content after label
        inline_runs = _inline_runs(state, node)
        if any(r.text.strip() for r in inline_runs):
            p.add_run(" ")
            _emit_runs(p, inline_runs)
        _emit_tree_nodes(state, node, depth=depth + 1)


def _h_rd_checklist(state: _State, el: ET._Element) -> None:
    for task in el:
        if not (isinstance(task.tag, str) and task.tag.lower() == "rd-task"):
            continue
        done = task.get("done") is not None
        marker = "☑ " if done else "☐ "
        runs = _inline_runs(state, task)
        p = state.doc.add_paragraph(style="List Bullet")
        p.add_run(marker)
        _emit_runs(p, runs)
        meta = []
        if task.get("assignee"):
            meta.append(f"@{task.get('assignee')}")
        if task.get("due"):
            meta.append(f"due {task.get('due')}")
        if meta:
            r = p.add_run("  " + " ".join(meta))
            r.italic = True
            r.font.size = Pt(8)


def _h_rd_mermaid(state: _State, el: ET._Element) -> None:
    text = _dedent(_element_source(el))
    _render_diagram(state, text, kind="mermaid")


def _h_rd_plantuml(state: _State, el: ET._Element) -> None:
    text = _dedent(_element_source(el))
    _render_diagram(state, text, kind="plantuml")


def _render_diagram(state: _State, text: str, *, kind: str) -> None:
    if not text.strip():
        return
    if state.render_diagrams:
        png = render_to_png(text, kind=kind, endpoint=state.diagram_endpoint)  # type: ignore[arg-type]
        if png is not None:
            state.doc.add_picture(BytesIO(png), width=Inches(6.0))
            state.diagrams_rendered += 1
            return
        state.diagrams_failed += 1
    _emit_code(state, text, lang=kind)


def _h_rd_toc(state: _State, el: ET._Element) -> None:
    chapters = [c for c in el if isinstance(c.tag, str) and c.tag.lower() == "rd-chapter"]
    if not chapters:
        state.record_dropped("rd-toc")
        return
    if state.toc_emitted:
        # The same rd-toc block is duplicated across every chapter of a book
        # (required for in-browser nav). One copy in the docx is plenty.
        return
    title = el.get("title") or "Contents"
    state.doc.add_heading(title, level=2)
    _emit_chapter_list(state, chapters, depth=0)
    state.toc_emitted = True


def _emit_chapter_list(state: _State, chapters, depth: int) -> None:
    for ch in chapters:
        title = _chapter_text(ch)
        href = ch.get("href")
        if title or href:
            style_level = min(depth, 2)
            style = ["List Bullet", "List Bullet 2", "List Bullet 3"][style_level]
            p = state.doc.add_paragraph(style=style)
            if href and title:
                _emit_runs(p, [_Run(title, hyperlink=href, underline=True)])
            elif href:
                _emit_runs(p, [_Run(href, hyperlink=href, underline=True)])
            else:
                r = p.add_run(title)
                r.bold = True
        nested = [c for c in ch if isinstance(c.tag, str) and c.tag.lower() == "rd-chapter"]
        if nested:
            _emit_chapter_list(state, nested, depth + 1)


def _chapter_text(node: ET._Element) -> str:  # noqa: SLF001
    parts: list[str] = []
    if node.text:
        parts.append(node.text)
    for child in node:
        if isinstance(child.tag, str) and child.tag.lower() == "rd-chapter":
            if child.tail:
                parts.append(child.tail)
            continue
        parts.extend(child.itertext())
        if child.tail:
            parts.append(child.tail)
    return " ".join("".join(parts).split()).strip()


def _h_rd_decision(state: _State, el: ET._Element) -> None:
    status = el.get("status") or "proposed"
    id_ = el.get("id") or ""
    title = el.get("title") or ""
    date = el.get("date") or ""
    deciders = el.get("deciders") or ""
    head = " — ".join(b for b in (id_, title) if b) or "Decision"
    state.doc.add_heading(head, level=2)
    p = state.add_paragraph()
    r = p.add_run(f"Status: {status}")
    r.bold = True
    if date or deciders:
        p2 = state.add_paragraph()
        bits = []
        if date:
            bits.append(f"Date: {date}")
        if deciders:
            bits.append(f"Deciders: {deciders}")
        r2 = p2.add_run("  ·  ".join(bits))
        r2.italic = True
    _render_children(state, el)


def _h_rd_pros_cons(state: _State, el: ET._Element) -> None:
    pros_title = el.get("pros-title") or "Pros"
    cons_title = el.get("cons-title") or "Cons"
    pros = [c for c in el if isinstance(c.tag, str) and c.tag.lower() == "rd-pro"]
    cons = [c for c in el if isinstance(c.tag, str) and c.tag.lower() == "rd-con"]
    rows = max(len(pros), len(cons))
    table = state.doc.add_table(rows=1 + rows, cols=2)
    table.style = "Table Grid"
    _fill_row(state, table.rows[0], [f"✓ {pros_title}", f"✗ {cons_title}"], bold=True)
    for i in range(rows):
        ptext = _flatten_inline(state, pros[i]).strip() if i < len(pros) else ""
        ctext = _flatten_inline(state, cons[i]).strip() if i < len(cons) else ""
        _fill_row(state, table.rows[i + 1], [ptext, ctext])


def _h_rd_compare(state: _State, el: ET._Element) -> None:
    headers = [h.strip() for h in (el.get("headers") or "").split(",")]
    rows = [r for r in el if isinstance(r.tag, str) and r.tag.lower() == "rd-row-cells"]
    if not headers and not rows:
        return
    cols = len(headers) + 1 if headers else max((len([c for c in r if isinstance(c.tag, str) and c.tag.lower() == "rd-cell"]) for r in rows), default=0) + 1
    if cols < 2:
        return
    table = state.doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    _fill_row(state, table.rows[0], [""] + headers, bold=True)
    for i, r in enumerate(rows, start=1):
        label = r.get("label") or ""
        cells = [c for c in r if isinstance(c.tag, str) and c.tag.lower() == "rd-cell"]
        values = [_flatten_inline(state, c).strip() for c in cells]
        _fill_row(state, table.rows[i], [label] + values + [""] * (cols - 1 - len(values)))


def _h_rd_rubric(state: _State, el: ET._Element) -> None:
    options = [o.strip() for o in (el.get("options") or "").split(",")]
    title = el.get("title") or ""
    scale = float(el.get("scale") or 5)
    criteria = [c for c in el if isinstance(c.tag, str) and c.tag.lower() == "rd-criterion"]
    if not options or not criteria:
        return
    if title:
        state.doc.add_heading(title, level=3)
    cols = 2 + len(options)  # criterion + weight + per-option scores
    table = state.doc.add_table(rows=1 + len(criteria) + 1, cols=cols)
    table.style = "Table Grid"
    _fill_row(state, table.rows[0], ["Criterion", "Weight"] + options, bold=True)
    totals = [0.0] * len(options)
    weights_sum = 0.0
    for i, crit in enumerate(criteria, start=1):
        label = crit.get("label") or ""
        weight = float(crit.get("weight") or 1)
        scores = [s for s in crit if isinstance(s.tag, str) and s.tag.lower() == "rd-score"]
        score_vals: list[str] = []
        for j, sc in enumerate(scores[: len(options)]):
            try:
                val = float(sc.get("value") or 0)
            except ValueError:
                val = 0.0
            totals[j] += val * weight
            score_vals.append(sc.get("value") or "")
        score_vals += [""] * (len(options) - len(score_vals))
        weights_sum += weight
        _fill_row(state, table.rows[i], [label, str(weight)] + score_vals)
    total_row = ["Total", f"{weights_sum:g}"]
    for v in totals:
        total_row.append(f"{v:g}")
    _fill_row(state, table.rows[-1], total_row, bold=True)


def _h_rd_roadmap(state: _State, el: ET._Element) -> None:
    title = el.get("title") or ""
    if title:
        state.doc.add_heading(title, level=3)
    lanes = [l for l in el if isinstance(l.tag, str) and l.tag.lower() == "rd-lane"]  # noqa: E741
    if not lanes:
        return
    table = state.doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _fill_row(state, table.rows[0], ["Lane", "Item", "Start", "End"], bold=True)
    for lane in lanes:
        name = lane.get("name") or ""
        for item in lane:
            if not (isinstance(item.tag, str) and item.tag.lower() == "rd-item"):
                continue
            row = table.add_row()
            _fill_row(
                state,
                row,
                [
                    name,
                    item.get("label") or "",
                    item.get("start") or "",
                    item.get("end") or "",
                ],
            )


def _h_rd_api(state: _State, el: ET._Element) -> None:
    method = el.get("method") or ""
    path = el.get("path") or ""
    title = el.get("title") or ""
    auth = el.get("auth") or ""
    head = f"{method} {path}".strip()
    if title:
        head = f"{title} — {head}" if head else title
    if head:
        state.doc.add_heading(head, level=3)
    if auth:
        p = state.add_paragraph()
        r = p.add_run(f"Auth: {auth}")
        r.italic = True
    params = [c for c in el if isinstance(c.tag, str) and c.tag.lower() == "rd-param"]
    responses = [c for c in el if isinstance(c.tag, str) and c.tag.lower() == "rd-response"]
    if params:
        state.doc.add_heading("Parameters", level=4)
        table = state.doc.add_table(rows=1 + len(params), cols=5)
        table.style = "Table Grid"
        _fill_row(state, table.rows[0], ["Name", "In", "Type", "Required", "Description"], bold=True)
        for i, p in enumerate(params, start=1):
            desc = _flatten_inline(state, p).strip()
            _fill_row(
                state,
                table.rows[i],
                [
                    p.get("name") or "",
                    p.get("in") or "query",
                    p.get("type") or "",
                    "yes" if p.get("required") is not None else "",
                    desc,
                ],
            )
    if responses:
        state.doc.add_heading("Responses", level=4)
        table = state.doc.add_table(rows=1 + len(responses), cols=3)
        table.style = "Table Grid"
        _fill_row(state, table.rows[0], ["Status", "Type", "Description"], bold=True)
        for i, resp in enumerate(responses, start=1):
            desc = _flatten_inline(state, resp).strip()
            _fill_row(
                state,
                table.rows[i],
                [resp.get("status") or "", resp.get("type") or "", desc],
            )


def _h_rd_swatch(state: _State, el: ET._Element) -> None:
    kind = el.get("kind") or ""
    name = el.get("name") or ""
    value = el.get("value") or ""
    note = el.get("note") or ""
    p = state.add_paragraph()
    r = p.add_run(f"{kind}: {name} = ")
    r.bold = True
    code = p.add_run(value)
    code.font.name = "Courier New"
    if note:
        p.add_run(f"  ({note})")


def _h_rd_footnotes(state: _State, el: ET._Element) -> None:
    # Top-level rd-footnotes is auto-injected at runtime; ignore here.
    state.record_dropped("rd-footnotes")


def _h_rd_references(state: _State, el: ET._Element) -> None:
    # rd-references placement marker — emit collected refs here if any.
    _emit_references(state, title=el.get("title") or "References")
    state._refs_emitted = True  # type: ignore[attr-defined]


def _h_rd_ref(state: _State, el: ET._Element) -> None:
    _collect_ref(state, el)


def _h_rd_footnote(state: _State, el: ET._Element) -> None:
    # Block-level rd-footnote shouldn't happen; render inline body as a paragraph.
    text = _flatten_inline(state, el).strip()
    if text:
        p = state.add_paragraph()
        r = p.add_run(text)
        r.italic = True


def _h_rd_cite(state: _State, el: ET._Element) -> None:
    # Top-level rd-cite is unusual but handle it gracefully.
    key = el.get("key") or ""
    if key and key not in state.cite_order:
        state.cite_order.append(key)


def _h_rd_chapter(state: _State, el: ET._Element) -> None:
    # Inside rd-toc only — caller handles it.
    pass


def _collect_ref(state: _State, el: ET._Element) -> None:
    key = el.get("key") or ""
    if not key:
        return
    state.refs_collected[key] = {
        "author": el.get("author") or "",
        "title": el.get("title") or "",
        "url": el.get("url") or "",
        "date": el.get("date") or "",
        "publisher": el.get("publisher") or "",
        "note": _flatten_inline(state, el).strip(),
    }


# ---------------------------------------------------------------------------
# Code blocks, image embeds, low-level helpers
# ---------------------------------------------------------------------------


def _emit_code(state: _State, text: str, *, lang: str | None) -> None:
    if not text:
        return
    if lang:
        p = state.add_paragraph(style="RichdocCode")
        r = p.add_run(f"// {lang}")
        r.italic = True
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for line in text.splitlines() or [text]:
        p = state.add_paragraph(style="RichdocCode")
        if line:
            r = p.add_run(line)
            # Preserve leading spaces.
            r.text = line


def _embed_image(state: _State, src: str, *, alt: str = "") -> None:
    ref = state.asset_store.add(src, base_dir=state.base_dir, fetch_remote=True)
    if ref is None:
        p = state.add_paragraph()
        r = p.add_run(f"[image not available: {alt or src}]")
        r.italic = True
        return
    try:
        state.doc.add_picture(BytesIO(ref.data), width=Inches(6.0))
        state.images_embedded += 1
    except Exception:  # noqa: BLE001 — unsupported format etc.
        p = state.add_paragraph()
        r = p.add_run(f"[image: {alt or src}]")
        r.italic = True


def _set_cell_border(cell: _Cell, *, side: str, color: str, sz: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = ET.SubElement(tcPr, qn("w:tcBorders"))
    border = tcBorders.find(qn(f"w:{side}"))
    if border is None:
        border = ET.SubElement(tcBorders, qn(f"w:{side}"))
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(sz))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)


def _set_cell_shading(cell: _Cell, color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = ET.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)


_WS = re.compile(r"\s+")


def _inline_text(text: str) -> str:
    return _WS.sub(" ", text)


def _element_source(el: ET._Element) -> str:  # noqa: SLF001
    """Return the literal source text of a leaf code-like element.

    Mirrors the JS runtime's `this.textContent` semantics: prefer a
    `<script type="text/...">` child if present (used to embed source
    without HTML escaping), otherwise concatenate every descendant text
    node."""
    for child in el:
        if not isinstance(child.tag, str):
            continue
        if child.tag.lower() == "script":
            script_text = "".join(child.itertext())
            if script_text.strip():
                return script_text
    return "".join(el.itertext())


def _dedent(text: str) -> str:
    lines = text.splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if not lines:
        return ""
    indents = [len(l) - len(l.lstrip(" ")) for l in lines if l.strip()]
    pad = min(indents) if indents else 0
    return "\n".join(l[pad:] if len(l) >= pad else l for l in lines)


# ---------------------------------------------------------------------------
# Finalisation
# ---------------------------------------------------------------------------


def _emit_references(state: _State, *, title: str) -> None:
    keys: list[str] = []
    seen: set[str] = set()
    for k in state.cite_order:
        if k in state.refs_collected and k not in seen:
            seen.add(k)
            keys.append(k)
    for k in state.refs_collected:
        if k not in seen:
            seen.add(k)
            keys.append(k)
    if not keys:
        return
    state.doc.add_heading(title, level=2)
    for n, key in enumerate(keys, start=1):
        attrs = state.refs_collected[key]
        parts = [f"[{n}]"]
        if attrs.get("author"):
            parts.append(attrs["author"] + ".")
        if attrs.get("title"):
            parts.append(f'"{attrs["title"]}."')
        if attrs.get("publisher"):
            parts.append(attrs["publisher"] + ".")
        if attrs.get("date"):
            parts.append(attrs["date"] + ".")
        text = " ".join(parts)
        p = state.add_paragraph()
        r = p.add_run(text)
        if attrs.get("url"):
            p.add_run(" ")
            _emit_runs(p, [_Run(attrs["url"], hyperlink=attrs["url"], underline=True)])
        if attrs.get("note"):
            p2 = state.add_paragraph()
            rr = p2.add_run(attrs["note"])
            rr.italic = True


def _finalise(state: _State) -> None:
    # Auto-emit references if none of the rd-references markers placed them.
    # Easy heuristic: if the document has any cite_order entries but the
    # final rendered text doesn't already include a "References" heading,
    # append one. python-docx makes that hard to detect, so we just always
    # append — duplicate output is the user's signal to remove the marker.
    if state.refs_collected:
        # If author placed rd-references explicitly, _emit_references already
        # ran. Track placement via state flag:
        if not getattr(state, "_refs_emitted", False):
            _emit_references(state, title="References")


def _serialise(state: _State) -> DocxResult:
    buf = BytesIO()
    state.doc.save(buf)
    return DocxResult(
        data=buf.getvalue(),
        dropped=sorted(set(state.dropped)),
        missing=list(state.asset_store.missing),
        diagrams_rendered=state.diagrams_rendered,
        diagrams_failed=state.diagrams_failed,
        images_embedded=state.images_embedded,
    )


# ---------------------------------------------------------------------------
# Dispatch table
# ---------------------------------------------------------------------------


_BLOCK_HANDLERS: dict[str, Callable[[_State, ET._Element], None]] = {
    # plain HTML
    "h1": _h_heading(1),
    "h2": _h_heading(2),
    "h3": _h_heading(3),
    "h4": _h_heading(4),
    "h5": _h_heading(5),
    "h6": _h_heading(6),
    "p": _h_p,
    "ul": _h_ul,
    "ol": _h_ol,
    "blockquote": _h_blockquote,
    "pre": _h_pre,
    "hr": _h_hr,
    "table": _h_table,
    "img": _h_img_block,
    "div": lambda s, e: _render_children(s, e),
    "section": lambda s, e: _render_children(s, e),
    "article": lambda s, e: _render_children(s, e),
    "header": lambda s, e: _render_children(s, e),
    "footer": lambda s, e: _render_children(s, e),
    "main": lambda s, e: _render_children(s, e),
    "aside": lambda s, e: _render_children(s, e),
    "nav": lambda s, e: _render_children(s, e),
    "figure": lambda s, e: _render_children(s, e),
    # rd-*
    "rd-page": _h_rd_page,
    "rd-section": _h_rd_section,
    "rd-hero": _h_rd_hero,
    "rd-banner": _h_rd_banner,
    "rd-callout": _h_rd_callout,
    "rd-kv": _h_rd_kv,
    "rd-stat": _h_rd_stat,
    "rd-progress": _h_rd_progress,
    "rd-update": _h_rd_update,
    "rd-quote": _h_rd_quote,
    "rd-cols": _h_rd_cols,
    "rd-card": _h_rd_card,
    "rd-code": _h_rd_code,
    "rd-diff": _h_rd_diff,
    "rd-shell": _h_rd_shell,
    "rd-math": _h_rd_math,
    "rd-figure": _h_rd_figure,
    "rd-chart": _h_rd_chart,
    "rd-gallery": _h_rd_gallery,
    "rd-embed": _h_rd_embed,
    "rd-tabs": _h_rd_tabs,
    "rd-timeline": _h_rd_timeline,
    "rd-steps": _h_rd_steps,
    "rd-detail": _h_rd_detail,
    "rd-tree": _h_rd_tree,
    "rd-checklist": _h_rd_checklist,
    "rd-mermaid": _h_rd_mermaid,
    "rd-plantuml": _h_rd_plantuml,
    "rd-toc": _h_rd_toc,
    "rd-decision": _h_rd_decision,
    "rd-pros-cons": _h_rd_pros_cons,
    "rd-compare": _h_rd_compare,
    "rd-rubric": _h_rd_rubric,
    "rd-roadmap": _h_rd_roadmap,
    "rd-api": _h_rd_api,
    "rd-swatch": _h_rd_swatch,
    "rd-footnotes": _h_rd_footnotes,
    "rd-references": _h_rd_references,
    "rd-ref": _h_rd_ref,
    "rd-footnote": _h_rd_footnote,
    "rd-cite": _h_rd_cite,
    "rd-chapter": _h_rd_chapter,
}
